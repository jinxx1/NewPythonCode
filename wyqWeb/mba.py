# -*- coding: utf-8 -*-
from docx import Document
import pprint
from pydocx import PyDocX
from bs4 import BeautifulSoup
import base64,json
import os.path
from pdfminer.layout import *
from wordall import WORDALL
from nltk.tokenize import word_tokenize
import sys,re,collections,nltk
from nltk.stem.wordnet import WordNetLemmatizer

from mysql import *

def requests_doxue():
    import requests, json, re

    # url = "https://ktiku.doxue.com/ktiku/examZH.html?5dfed387321ee65e2f69a2f2"
    url = "https://ktiku.doxue.com/port/paper/getPaper?callback=jQuery19107162448831899504_1602602241057&uid=0&paperID=5dfed387321ee65e2f69a2f2&record_id=&_=1602602241058"
    brow = requests.get(url=url)
    Html = brow.text.encode(brow.encoding).decode('utf-8')
    # print(Html)
    Html1 = re.findall("\{.*\}", Html)[0]

    jsonT = json.loads(Html1)['data']['paper']
    # print(jsonT.keys())
    Title = jsonT['paper']
    filePathName = '_'.join([str(jsonT['publish_year']), jsonT['subject']['subject_code']]) + ".html"
    print(Title)
    print(filePathName)

    for i in jsonT['questions']:
        for n in i:
            print(n)
            print('*************************')
            # break
        print('--------------------------------')

def file_name_walk(file_dir):
    listpath = []
    for root, dirs, files in os.walk(file_dir):
        dictTemp = {}
        dictTemp['root'] = root# 当前目录路径
        dictTemp['dirs'] = dirs# 当前路径下所有子目录
        dictTemp['files'] = files# 当前路径下所有非目录子文件
        listpath.append(dictTemp)
    return listpath
null = None

wordFilePath = r"D:/mba_exam/在职MBA管理/在职管理类联考历年真题/2009.docx"

PDFpath = r"D:/mba_exam/在职MBA管理/在职管理类联考历年真题/2009.pdf"

def pdfminerprocess(pdfpath):
    # 导入库
    from pdfminer.pdfparser import PDFParser
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfpage import PDFTextExtractionNotAllowed
    from pdfminer.pdfinterp import PDFResourceManager
    from pdfminer.pdfinterp import PDFPageInterpreter
    from pdfminer.pdfdevice import PDFDevice

    from pdfminer.converter import PDFPageAggregator

    # 提供初始密码
    password = ''
    # 没有密码可以初始密码
    # document.initialize()

    # 打开pdf文件
    fp = open(pdfpath, 'rb')

    # 从文件句柄创建一个pdf解析对象
    parser = PDFParser(fp)

    # 创建pdf文档对象，存储文档结构
    document = PDFDocument(parser, password)

    # 创建一个pdf资源管理对象，存储共享资源
    rsrcmgr = PDFResourceManager()

    laparams = LAParams()

    # 创建一个device对象
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)

    # 创建一个解释对象
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    # 处理包含在文档中的每一页
    llist1 = []

    for page in PDFPage.create_pages(document):
        interpreter.process_page(page)
        layout = device.get_result()
        for x in layout:
            # 获取文本对象
            if isinstance(x, LTTextBox):
                atext = x.get_text().strip()
                if len(atext) <3:
                    continue
                llist1.append(reSub(atext))
            # 获取图片对象
            if isinstance(x, LTImage):
                print('这里获取到一张图片')
            # 获取 figure 对象
            if isinstance(x, LTFigure):
                for nx in x:
                    import pdir
                    print(nx)
                    print(nx.imagemask)
                    print(pdir(nx))
                print('这里获取到一个 figure 对象')
    return ''.join(llist1)

def nltk_word(strword):
    import nltk
    value = nltk.sent_tokenize(strword)
    llistw = []
    for i in value:
        words = nltk.word_tokenize(text=i)
        if len(words) > 0:
            for nn in words:
                s = re.sub("[A-Z]$", '', nn)
                if len(s)<3:
                    continue
                if s:
                    llistw.append(s)
    return llistw

def jsonload(jsonPath):
    try:
        with open(jsonPath, 'r') as jf:
            jsonLoad = json.load(jf)
            jf.close()
    except:
        with open(jsonPath, 'w') as jf:
            jsonLoad = {}
            jf.close()
    return jsonLoad

def jsonwrite(jsonPath,item):
    with open(jsonPath, 'w') as jf:
        json.dump(item,jf)
        jf.close()

def precessWord(wordlist):
    ddict = jsonload(jsonFile)
    for i in wordlist:
        if i not in ddict.keys():
            ddict[i] = 1
        else:
            ddict[i] += 1
    return ddict

def greall(path):
    document = Document(path)
    llistall = []
    llist = []
    for num,i in enumerate(document.paragraphs):
        strword = i.text.strip()
        if not strword:
            continue
        cutList_regx = re.findall("^List.*",strword)
        if cutList_regx:
            continue
        startword_regx = re.findall("^[a-zA-Z].*\[.*\]",strword)
        if startword_regx:
            llistall.append(llist)
            llist = []
        llist.append(strword)

        print('---------------------')
    print(llistall)

def doc_2_docx():
    from win32com import client as wc
    word_ob = wc.Dispatch('Word.Application')
    docPath = r"D:\mba_exam\英语二真题\英语二真题\真题集（纯真题，可直接打印）"
    greallpath = r"D:\PythonCode\wyqWeb\engexam"
    docPathall = file_name_walk(docPath)

    for n in docPathall:
        for nn in n['files']:
            docFile = os.path.join(n['root'],nn)
            if docFile.endswith(".doc"):
                asSaveFilePath = os.path.join(greallpath,nn)
                doc = word_ob.Documents.Open(docFile)
                doc.SaveAs("{}x".format(asSaveFilePath),12)
                doc.Close()
    word_ob.Quit()
    print('ok了')

def reSub(strword):
    import re

    pat_letter = re.compile(r'[^a-zA-Z \']+')
    new_text = pat_letter.sub(' ', strword).strip().lower()
    # to find the 's following the pronouns. re.I is refers to ignore case
    pat_is = re.compile("(it|he|she|that|this|there|here)(\'s)", re.I)
    # to find the 's following the letters
    pat_s = re.compile("(?<=[a-zA-Z])\'s")
    # to find the ' following the words ending by s
    pat_s2 = re.compile("(?<=s)\'s?")
    # to find the abbreviation of not
    pat_not = re.compile("(?<=[a-zA-Z])n\'t")
    # to find the abbreviation of would
    pat_would = re.compile("(?<=[a-zA-Z])\'d")
    # to find the abbreviation of will
    pat_will = re.compile("(?<=[a-zA-Z])\'ll")
    # to find the abbreviation of am
    pat_am = re.compile("(?<=[I|i])\'m")
    # to find the abbreviation of are
    pat_are = re.compile("(?<=[a-zA-Z])\'re")
    # to find the abbreviation of have
    pat_ve = re.compile("(?<=[a-zA-Z])\'ve")

    new_text = pat_is.sub(r"\1 is", new_text)
    new_text = pat_s.sub("", new_text)
    new_text = pat_s2.sub("", new_text)
    new_text = pat_not.sub(" not", new_text)
    new_text = pat_would.sub(" would", new_text)
    new_text = pat_will.sub(" will", new_text)
    new_text = pat_am.sub(" am", new_text)
    new_text = pat_are.sub(" are", new_text)
    new_text = pat_ve.sub(" have", new_text)
    new_text = new_text.replace('\'', ' ')

    return new_text


def get_wordnet_pos(treebank_tag):
    if treebank_tag.startswith('J'):
        return nltk.corpus.wordnet.ADJ
    elif treebank_tag.startswith('V'):
        return nltk.corpus.wordnet.VERB
    elif treebank_tag.startswith('N'):
        return nltk.corpus.wordnet.NOUN
    elif treebank_tag.startswith('R'):
        return nltk.corpus.wordnet.ADV
    else:
        return ''




if __name__ == '__main__':

    # from nltk.stem.porter import *
    # from nltk.stem.snowball import SnowballStemmer
    # import nltk
    # stemmer = SnowballStemmer("english")
    #
    #
    # wordList = ['depletion', 'epithet', 'confrontation', 'exemplify', 'justifiable', 'bearing', 'propel', 'renovate',
    #             'inconstancy', 'strip', 'primeval', 'cosmos', 'balloon', 'refinement', 'torrent']
    #
    # aa = stemmer.stem('confrontation')
    # bb = stemmer.stem(aa)
    # print(aa)
    # print(bb)
    # for i in wordList:
    #     aa = stemmer.stem(i)
    #     print(aa)

    # exit()

    lmtzr = WordNetLemmatizer()
    greallpath = r"D:\PythonCode\wyqWeb\engexam"
    docPathall = file_name_walk(greallpath)
    docxList = []
    for n in docPathall:
        for nn in n['files']:
            docFile = os.path.join(n['root'],nn)
            if docFile.endswith(".docx") and 'greall' not in docFile:
                docxList.append(docFile)
    ALL_WORD_LIST = []
    for num,docxFP in enumerate(docxList):
        pydocxFile = Document(docxFP)
        for para in pydocxFile.paragraphs:
            getWord = reSub(para.text)

            tag = nltk.pos_tag(word_tokenize(getWord))
            tag = [i[0] for i in tag if i and len(i[0])>2]
            if not tag:
                continue
            ALL_WORD_LIST.extend(tag)

    collcounts = collections.Counter(ALL_WORD_LIST)
    llist=[]
    for i in collcounts.items():
        ddict = {}
        ddict['word'] = i[0]
        ddict['showCount'] = i[1]
        llist.append(ddict)
    exam_word = pd.DataFrame(llist)
    all_greWord_df = get_gre_co_pandas()
    df = pd.merge(exam_word,all_greWord_df,how='inner',on='word')
    cutList = ['text','even','want','down','group','list','low','meet']
    df = df[~df['word'].isin(cutList)]
    df.sort_values('showCount',ascending=False,inplace=True)
    df.to_csv('examWord.csv',encoding='utf-8')
    df.to_excel('examWord.xlsx', encoding='utf-8')
    print(df)
    print(len(df))










