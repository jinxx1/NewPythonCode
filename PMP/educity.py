import pprint
import re

import requests
import json
import csv
from bs4 import BeautifulSoup
import pandas as pd

import sqlalchemy
import xlrd

MYSQLINFO = {
	"HOST": "localhost",
	"DBNAME": "pmp",
	"USER": "root",
	"PASSWORD": "040304",
	"PORT": 3306
}

conStr = 'mysql+pymysql://{USER}:{PASSWORD}@{HOST}:{PORT}/{DBNAME}?charset=utf8mb4'.format(USER=MYSQLINFO['USER'],
                                                                                           PASSWORD=MYSQLINFO[
	                                                                                           'PASSWORD'],
                                                                                           HOST=MYSQLINFO['HOST'],
                                                                                           PORT=MYSQLINFO['PORT'],
                                                                                           DBNAME=MYSQLINFO[
	                                                                                           'DBNAME'])

mysqlcon = sqlalchemy.create_engine(conStr)


def excel_process():
	import openpyxl
	xlsroot = r"D:\PythonCode\PMP\eductity1.xlsx"
	workbook = openpyxl.load_workbook(xlsroot)
	sheet_name = workbook.get_sheet_names()
	print(sheet_name)
	llist = []
	for sheetname in sheet_name:
		ddict = {}
		ddict['exam_title'] = sheetname
		ddict['data'] = []
		sheet = workbook.get_sheet_by_name(ddict['exam_title'])
		ddict['data'] = [item.value for item in list(sheet.columns)[0]]
		llist.append(ddict)
	return llist


def getContent(quest_id):
	global HEA
	item = {}
	url = "https://uc.educity.cn/ucapi/uc/paper/loadShitiInfo.do"
	postdate = {'id': str(quest_id)}
	brow = requests.post(url=url, data=postdate, headers=HEA)
	jsonT = json.loads(brow.text)['model']['shiti']
	item['quest_id'] = quest_id
	item['school'] = 'educity'
	item['title'] = str(jsonT['questionNum']) + "、" + jsonT['tiganDelHTMLTag']
	Map = [x for x in jsonT['questionMap']][0]
	item['A'] = Map['A']
	item['B'] = Map['B']
	item['C'] = Map['C']
	item['D'] = Map['D']
	item['quest_answers'] = jsonT['answerStr']
	item['quest_analysis'] = jsonT['analysis']
	item['sectionNames'] = jsonT['sectionNames']
	return item


def get_educity_quest():
	HEA = {
		"cookie": "JESONG_USER_ID=01000000013264186911601145863116; _sid_=2f62f0dea12fa0206a05b5a80bf5019b; Hm_lvt_555d9dcffdcb317595de82b0fc125cdf=1621869116,1622708749,1623054426,1623830671; _rme=T; cstk=493da130d698d232953684fec19df838; JESONG_VISITOR_ID=01000000013264383072533879479104; autoAlertNum_PC_13264=0; autoAlertNum_M_13264=0; uname=%u7504%u65B0; _subjectCode_=10031006; Hm_lpvt_555d9dcffdcb317595de82b0fc125cdf=1623832240",
		"user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.101 Safari/537.36",
	}
	for excel in excel_process():
		llist = []
		for quest_id in excel['data']:
			try:
				item = getContent(quest_id)
				item['exam_title'] = excel['exam_title']
			except:
				print('++++++++++++++++++++++', quest_id)
			print(item)
			print('--------------------')
			llist.append(item)
		df = pd.DataFrame(llist)
		df.to_sql(name='exam_copy1', con=mysqlcon, if_exists='append', index=False, chunksize=1000)

def filter_tags(htmlstr):
    # 先过滤CDATA
    re_cdata = re.compile('//<!\[CDATA\[[^>]*//\]\]>', re.I)  # 匹配CDATA
    re_script = re.compile('<\s*script[^>]*>[^<]*<\s*/\s*script\s*>', re.I)  # Script
    re_style = re.compile('<\s*style[^>]*>[^<]*<\s*/\s*style\s*>', re.I)  # style
    re_br = re.compile('<br\s*?/?>')  # 处理换行
    re_h = re.compile('</?\w+[^>]*>')  # HTML标签
    re_comment = re.compile('<!--[^>]*-->')  # HTML注释
    s = re_cdata.sub('', htmlstr)  # 去掉CDATA
    s = re_script.sub('', s)  # 去掉SCRIPT
    s = re_style.sub('', s)  # 去掉style
    s = re_br.sub('\n', s)  # 将br转换为换行
    s = re_h.sub('', s)  # 去掉HTML 标签
    s = re_comment.sub('', s)  # 去掉HTML注释
    # 去掉多余的空行
    blank_line = re.compile('\n+')
    s = blank_line.sub('\n', s)
    s = replaceCharEntity(s)  # 替换实体
    return s.replace('\n', '').replace('"', '').replace('\\', '').replace("   、 ", '、').replace("?n",
                                                                                                '?</h2><h2>').replace(
        "  ", '')

def replaceCharEntity(htmlstr):
    CHAR_ENTITIES = {'nbsp': ' ', '160': ' ',
                     'lt': '<', '60': '<',
                     'gt': '>', '62': '>',
                     'amp': '&', '38': '&',
                     'quot': '"', '34': '"', }

    re_charEntity = re.compile(r'&#?(?P<name>\w+);')
    sz = re_charEntity.search(htmlstr)
    while sz:
        entity = sz.group()  # entity全称，如&gt;
        key = sz.group('name')  # 去除&;后entity,如&gt;为gt
        try:
            htmlstr = re_charEntity.sub(CHAR_ENTITIES[key], htmlstr, 1)
            sz = re_charEntity.search(htmlstr)
        except KeyError:
            # 以空串代替
            htmlstr = re_charEntity.sub('', htmlstr, 1)
            sz = re_charEntity.search(htmlstr)
    return htmlstr

def htmlget(html1):

	soup = BeautifulSoup(html1, 'lxml')
	try:
		img = soup.find('img')
		img_html = '''<img src="{}">'''.format(img.get('src'))
		html_title = "<p>" + img_html + "</p>" + filter_tags(html1)
	except:
		html_title = filter_tags(html1)
	return html_title


def get_quest_form_db(limit=None, exam_title1=None):
	if not exam_title1:
		print('input exam_title')
		return None
	if limit:
		exc = "SELECT exam_title,quest_answers,title,A,B,C,D,quest_analysis,quest_id,sectionNames FROM exam_copy1 WHERE school = 'educity' and exam_title = '{}' limit {};".format(
			exam_title1, limit)
	else:
		exc = "SELECT exam_title,quest_answers,title,A,B,C,D,quest_analysis,quest_id,sectionNames FROM exam_copy1 WHERE school = 'educity' and exam_title = '{}';".format(
			exam_title1)
	allinfo = mysqlcon.execute(exc)
	llist = []
	for i in allinfo:
		ddict = {}
		ddict['exam_title'] = i[0]
		ddict['quest_answers'] = i[1]
		ddict['title'] = i[2].replace('1、','')
		ddict['A'] = i[3]
		ddict['B'] = i[4]
		ddict['C'] = i[5]
		ddict['D'] = i[6]
		ddict['quest_analysis'] = htmlget(i[7])
		ddict['quest_id'] = i[8]
		ddict['sectionNames'] = i[9]
		llist.append(ddict)
	return llist


def word_get():
	from docx import Document
	from docx.shared import Pt
	from docx.oxml.ns import qn
	import os
	path = r"D:\PythonCode\PMP\educityExam"
	sheetName = ['第九套', '第八套', '第七套', '第六套', '第五套', '第四套', '第三套', '第二套', '第一套', '期中']
	for examTitle in sheetName:
		allinfo = get_quest_form_db(limit=None, exam_title1=examTitle)
		document = Document()
		keyWord = examTitle
		document.add_heading('2021模拟考试（{}） '.format(examTitle))
		exam_path = os.path.join(path, keyWord + '.docx')
		for item in allinfo:
			document.add_heading(item['title'], 1)
			corrt = ['A、' + item['A'], 'B、' + item['B'], 'C、' + item['C'], 'D、' + item['D']]
			corrt_str = '\n'.join(corrt)
			paragraph = document.add_paragraph()
			run = paragraph.add_run(corrt_str)
			run.font.size = Pt(10)
			run.font.name = u'宋体'
			r = run._element
			r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

			answer = ['正确答案：{}'.format(item['quest_answers']), '问题解析：{}'.format(item['quest_analysis']),'对应知识点：{}'.format(item['sectionNames'])]
			answer_str = '\n'.join(answer)
			paragraph = document.add_paragraph()
			run = paragraph.add_run(answer_str)
			run.font.size = Pt(8)
			run.font.name = u'宋体'
			run.italic = True
			r = run._element
			r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

		document.save(exam_path)


if __name__ == '__main__':
	word_get()

	exit()
	llist = get_quest_form_db(limit=10, exam_title1='第一套')

	pprint.pprint(llist)
